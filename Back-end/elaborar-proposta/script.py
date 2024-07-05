from docx import Document
import os
import win32com.client

def substituir_campos(doc, campos):
    for paragrafo in doc.paragraphs:
        for campo, valor in campos.items():
            if campo in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(campo, valor)

    for tabela in doc.tables:
        for row in tabela.rows:
            for cell in row.cells:
                for campo, valor in campos.items():
                    if campo in cell.text:
                        cell.text = cell.text.replace(campo, valor)

def salvar_em_pdf(doc, caminho_pdf):
    doc.save(caminho_pdf)

def converter_para_pdf(docx_path, pdf_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    doc.SaveAs(pdf_path, FileFormat=17)  # 17 é o código de formato para PDF
    doc.Close()
    word.Quit()

def main():
    # Obtendo o caminho do script atual
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Construindo o caminho para o documento Word usando os.path.join
    doc_path = os.path.join(script_dir, 'modelo_base.docx')

    # Carregando o documento Word
    documento = Document(doc_path)

    # Definindo os campos que você deseja substituir
    dados_para_substituir = {
        '{{NOME}}': 'Ane Karoline Lima',
        '{{DATA}}': '06',
        '{{MES}}': 'maio',
        '{{VENDEDOR}}': 'Danilo França'
        # Adicione mais campos conforme necessário
    }

    # Substituindo os campos no documento
    substituir_campos(documento, dados_para_substituir)

    # Construindo o caminho para o novo documento Word
    novo_doc_path = os.path.join(script_dir, 'Proposta Comercial QRPOINT.docx')

    # Salvando o documento modificado em formato Word
    salvar_em_pdf(documento, novo_doc_path)

    # Construindo o caminho para o novo documento PDF
    novo_pdf_path = os.path.join(script_dir, 'Proposta Comercial QRPOINT.pdf')

    # Convertendo o documento Word para PDF
    converter_para_pdf(novo_doc_path, novo_pdf_path)

if __name__ == "__main__":
    main()
